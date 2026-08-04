"""
docx_export_service.py — 生成 Word 文档（真脚注 / 参考书目）

真脚注实现思路：python-docx 不支持原生脚注对象，因此先用 Jinja2 渲染一份
Flat ODT（纯 XML，OASIS 规范原生支持 <text:note> 脚注），再用容器内置的
LibreOffice headless 把它转换成 .docx——转换出来的脚注是 Word 认可的真脚注
对象（可在 Word 里正常增删、自动重新编号）。

若运行环境没有装 LibreOffice（例如本地开发机），自动降级为用 python-docx
生成一份"编号脚注草稿"文档，保证功能始终可用，只是脚注编号需要用户在贴入
正式论文时手动核对。
"""

import shutil
import subprocess
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

import storage
from services.binaries import soffice_bin
from services.citation_service import BibliographyItem, RenderedFootnote

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"

_env = Environment(
    loader=FileSystemLoader(str(TEMPLATES_DIR)),
    autoescape=select_autoescape(["j2"]),
)

DEFAULT_FONT = "PMingLiU"


def soffice_available() -> bool:
    return soffice_bin() is not None


def soffice_bin_path() -> str | None:
    return soffice_bin()


def _soffice_bin() -> str:
    return soffice_bin() or "soffice"


def _set_default_font(document: Document, font_name: str = DEFAULT_FONT) -> None:
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(12)
    # 中文字体在 python-docx 里需要额外设置东亚字体名，否则 Word 显示仍是默认字体
    rpr = style.element.get_or_add_rPr()
    from docx.oxml.ns import qn

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _fallback_footnote_docx(
    items: list[RenderedFootnote], dest_path: str, doc_title: str, project_name: str
) -> str:
    document = Document()
    _set_default_font(document)
    document.add_heading(doc_title, level=1)
    document.add_paragraph(
        f"以下为「{project_name}」引用篮按收集顺序生成的脚注草稿"
        "（未检测到 LibreOffice，暂以编号列表呈现，可手动复制进 Word 脚注）。"
    )
    for item in items:
        p = document.add_paragraph()
        p.add_run(f"[{item.order}] ").bold = True
        p.add_run(item.text)
        if item.quoted_text:
            quote_p = document.add_paragraph()
            quote_p.add_run(f"原文摘录：{item.quoted_text}").italic = True
    document.save(dest_path)
    return dest_path


def render_footnote_docx(
    items: list[RenderedFootnote],
    dest_path: str,
    doc_title: str = "引用脚注导出",
    project_name: str = "引用篮",
) -> str:
    if not soffice_available():
        return _fallback_footnote_docx(items, dest_path, doc_title, project_name)

    template = _env.get_template("footnotes.fodt.j2")
    render_items = [
        {
            "order": item.order,
            "footnote_text": item.text,
            "quoted_excerpt": (item.quoted_text or "")[:80],
        }
        for item in items
    ]
    xml_content = template.render(doc_title=doc_title, project_name=project_name, items=render_items)

    work_dir = storage.EXPORTS_DIR / uuid.uuid4().hex
    work_dir.mkdir(parents=True, exist_ok=True)
    fodt_path = work_dir / "footnotes.fodt"
    fodt_path.write_text(xml_content, encoding="utf-8")

    try:
        subprocess.run(
            [
                _soffice_bin(),
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(work_dir),
                str(fodt_path),
            ],
            check=True,
            capture_output=True,
            timeout=120,
        )
        converted = work_dir / "footnotes.docx"
        if converted.exists():
            shutil.move(str(converted), dest_path)
            shutil.rmtree(work_dir, ignore_errors=True)
            return dest_path
    except Exception:
        pass

    shutil.rmtree(work_dir, ignore_errors=True)
    return _fallback_footnote_docx(items, dest_path, doc_title, project_name)


def render_bibliography_docx(
    items: list[BibliographyItem], dest_path: str, doc_title: str = "参考书目"
) -> str:
    document = Document()
    _set_default_font(document)
    document.add_heading(doc_title, level=1)

    needs_manual_check = any(not it.stroke_estimated for it in items)
    if needs_manual_check:
        note = document.add_paragraph()
        note.add_run(
            "提示：部分作者姓氏未匹配到内置笔画表，已临时按拼音顺序排列，"
            "建议核对后手动调整顺序。"
        ).italic = True

    for item in items:
        document.add_paragraph(item.text)

    document.save(dest_path)
    return dest_path
